"""
Stage 1: Convert documents (PDF, DOCX, etc.) to structured markdown.
"""
import os
from pathlib import Path
from typing import Dict, Optional, Any
from loguru import logger

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    logger.warning("PyMuPDF not available")
    PYMUPDF_AVAILABLE = False

try:
    from markitdown import MarkItDown
    MARKITDOWN_AVAILABLE = True
except ImportError:
    logger.warning("MarkItDown not available")
    MARKITDOWN_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    logger.warning("python-docx not available")
    PYTHON_DOCX_AVAILABLE = False


class MarkdownConverter:
    """Convert various document formats to markdown."""

    def __init__(self):
        """Initialize markdown converter."""
        self.markitdown = MarkItDown() if MARKITDOWN_AVAILABLE else None

    def convert(self, file_path: str) -> Dict[str, Any]:
        """
        Convert document to markdown.

        Args:
            file_path: Path to document file

        Returns:
            Dictionary with markdown content and metadata

        Raises:
            ValueError: If file format not supported
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        # Get file extension
        ext = file_path.suffix.lower()

        # Route to appropriate converter
        if ext == ".pdf":
            return self._convert_pdf(file_path)
        elif ext == ".docx":
            return self._convert_docx(file_path)
        elif ext == ".md":
            return self._read_markdown(file_path)
        elif ext in [".txt"]:
            return self._read_text(file_path)
        else:
            # Try markitdown as fallback
            if self.markitdown:
                return self._convert_with_markitdown(file_path)
            else:
                raise ValueError(f"Unsupported file format: {ext}")

    def _convert_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Convert PDF to markdown using PyMuPDF."""
        if not PYMUPDF_AVAILABLE:
            raise RuntimeError("PyMuPDF not installed. Install with: pip install pymupdf")

        logger.info(f"Converting PDF: {file_path.name}")

        doc = fitz.open(file_path)
        markdown_parts = []
        metadata = {
            "page_count": len(doc),
            "title": doc.metadata.get("title", ""),
            "author": doc.metadata.get("author", ""),
            "source_file": str(file_path),
        }

        # Common watermark patterns to filter out
        watermark_patterns = [
            r"(?i)watermark",
            r"(?i)draft",
            r"(?i)confidential",
            r"(?i)do not copy",
            r"(?i)internal use only",
            r"(?i)approved",
            r"_Approved",
        ]

        for page_num, page in enumerate(doc, start=1):
            # Extract text blocks
            blocks = page.get_text("dict")["blocks"]

            # Extract tables first and collect their bounding boxes
            table_bboxes = []
            tables = page.find_tables()
            if tables.tables:
                for table in tables:
                    markdown_table = self._extract_table_from_pdf(table)
                    if markdown_table:
                        markdown_parts.append(markdown_table)
                        # Store table bounding box to avoid converting table content to headings
                        table_bboxes.append(table.bbox)

            for block in blocks:
                if block["type"] == 0:  # Text block
                    for line in block["lines"]:
                        # Check if this line is inside a table bounding box
                        line_bbox = line.get("bbox", None)
                        in_table_area = False
                        if line_bbox:
                            # Check if line overlaps with any table bbox
                            for table_bbox in table_bboxes:
                                # Check overlap: line is inside table if its position is within table bounds
                                if (line_bbox[0] >= table_bbox[0] - 5 and
                                    line_bbox[2] <= table_bbox[2] + 5 and
                                    line_bbox[1] >= table_bbox[1] - 5 and
                                    line_bbox[3] <= table_bbox[3] + 5):
                                    in_table_area = True
                                    break

                        # Combine spans into line text
                        line_text = ""
                        font_sizes = []

                        for span in line["spans"]:
                            text = span["text"].strip()
                            if text:
                                # Skip watermark text
                                if self._is_watermark(text, watermark_patterns):
                                    continue

                                font_sizes.append(span["size"])
                                line_text += text + " "

                        line_text = line_text.strip()

                        if line_text:
                            # Skip heading conversion for text inside table areas
                            if in_table_area:
                                markdown_parts.append(line_text)
                            else:
                                # Determine heading level based on font size
                                avg_font_size = sum(font_sizes) / len(font_sizes) if font_sizes else 0

                                if avg_font_size > 16:
                                    markdown_parts.append(f"\n# {line_text}\n")
                                elif avg_font_size > 14:
                                    markdown_parts.append(f"\n## {line_text}\n")
                                elif avg_font_size > 12:
                                    markdown_parts.append(f"\n### {line_text}\n")
                                else:
                                    markdown_parts.append(line_text)

                elif block["type"] == 1:  # Image block
                    # Note image presence
                    markdown_parts.append(f"\n[Image on page {page_num}]\n")

            # Add page break marker
            markdown_parts.append(f"\n<!-- Page {page_num} -->\n")

        doc.close()

        markdown_content = "\n".join(markdown_parts)

        # Clean up excessive newlines (formatting only, no content removal)
        markdown_content = self._clean_markdown(markdown_content)

        # Note: Orphaned table content lines are now handled by skipping heading conversion
        # for text inside table bounding boxes. No need to remove them separately.

        # Note: Watermark, header, footer removal should be handled by /documents/cleanfile API
        # This markdown converter is responsible ONLY for text extraction and formatting

        return {
            "markdown": markdown_content,
            "metadata": metadata,
        }

    def _is_watermark(self, text: str, patterns: list) -> bool:
        """Check if text is likely a watermark."""
        import re

        # Check against watermark patterns
        for pattern in patterns:
            if re.search(pattern, text):
                return True

        # Check if text is very small or very large (common watermark sizes)
        # This would need font size info from caller

        # Check if text is semi-transparent or rotated (common watermark attributes)
        # This would need additional metadata from PDF

        return False

    def _extract_table_from_pdf(self, table) -> str:
        """Extract table from PDF using PyMuPDF table detection."""
        try:
            # Get table data
            table_data = table.extract()

            if not table_data or len(table_data) < 2:
                return ""

            # Build markdown table
            markdown_rows = []

            for i, row in enumerate(table_data):
                # Clean cells
                cells = [str(cell).strip() if cell else "" for cell in row]

                # Skip empty rows
                if not any(cells):
                    continue

                markdown_row = "| " + " | ".join(cells) + " |"
                markdown_rows.append(markdown_row)

                # Add separator after first row (header)
                if i == 0:
                    separator = "| " + " | ".join(["---"] * len(cells)) + " |"
                    markdown_rows.append(separator)

            if len(markdown_rows) > 2:  # At least header + separator + 1 row
                return "\n" + "\n".join(markdown_rows) + "\n"

            return ""
        except Exception as e:
            logger.warning(f"Failed to extract table: {e}")
            return ""

    def _convert_docx(self, file_path: Path) -> Dict[str, Any]:
        """Convert DOCX to markdown."""
        if not PYTHON_DOCX_AVAILABLE:
            raise RuntimeError("python-docx not installed. Install with: pip install python-docx")

        logger.info(f"Converting DOCX: {file_path.name}")

        doc = DocxDocument(file_path)
        markdown_parts = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Check style for headings
            style = para.style.name.lower()
            if "heading 1" in style:
                markdown_parts.append(f"\n# {text}\n")
            elif "heading 2" in style:
                markdown_parts.append(f"\n## {text}\n")
            elif "heading 3" in style:
                markdown_parts.append(f"\n### {text}\n")
            elif "heading" in style:
                markdown_parts.append(f"\n#### {text}\n")
            else:
                markdown_parts.append(text)

        # Extract tables
        for table in doc.tables:
            markdown_parts.append(self._table_to_markdown(table))

        markdown_content = "\n".join(markdown_parts)
        markdown_content = self._clean_markdown(markdown_content)

        metadata = {
            "source_file": str(file_path),
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
        }

        return {
            "markdown": markdown_content,
            "metadata": metadata,
        }

    def _table_to_markdown(self, table) -> str:
        """Convert DOCX table to markdown table."""
        markdown_rows = []

        for i, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            markdown_row = "| " + " | ".join(cells) + " |"
            markdown_rows.append(markdown_row)

            # Add separator after first row (header)
            if i == 0:
                separator = "| " + " | ".join(["---"] * len(cells)) + " |"
                markdown_rows.append(separator)

        return "\n" + "\n".join(markdown_rows) + "\n"

    def _read_markdown(self, file_path: Path) -> Dict[str, Any]:
        """Read existing markdown file."""
        logger.info(f"Reading markdown: {file_path.name}")

        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        metadata = {
            "source_file": str(file_path),
        }

        return {
            "markdown": content,
            "metadata": metadata,
        }

    def _read_text(self, file_path: Path) -> Dict[str, Any]:
        """Read plain text file."""
        logger.info(f"Reading text: {file_path.name}")

        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        metadata = {
            "source_file": str(file_path),
        }

        return {
            "markdown": content,
            "metadata": metadata,
        }

    def _convert_with_markitdown(self, file_path: Path) -> Dict[str, Any]:
        """Convert using MarkItDown library."""
        if not self.markitdown:
            raise RuntimeError("MarkItDown not available")

        logger.info(f"Converting with MarkItDown: {file_path.name}")

        result = self.markitdown.convert(str(file_path))

        metadata = {
            "source_file": str(file_path),
        }

        return {
            "markdown": result.text_content,
            "metadata": metadata,
        }

    def _remove_orphaned_table_lines(self, markdown: str) -> str:
        """
        Remove orphaned table content lines that appear after table headings.

        These are often duplicates or broken extractions from PDF tables that weren't
        properly detected as table structures. They appear as separate text lines
        instead of being part of the markdown table.
        """
        import re

        lines = markdown.split('\n')
        result = []
        i = 0

        # Common table content patterns that shouldn't appear as isolated lines
        table_content_words = {
            "container khô", "container lạnh", "tàu/sà lan", "phương án",
            "20'", "40'", "45'", "tt", "tàu/", "bãi", "xe  bãi",
            "giao/nhận", "loại container", "imdg", "oog"
        }

        # Also add variants with Unicode smart quotes (file might have U+2019 instead of U+0027)
        smart_quote_variants = []
        for word in list(table_content_words):
            if "'" in word:  # If word contains ASCII apostrophe (U+0027)
                # Add variant with Unicode right single quotation mark (U+2019)
                smart_word = word.replace("'", "\u2019")
                if smart_word not in table_content_words:
                    smart_quote_variants.append(smart_word)
        table_content_words.update(smart_quote_variants)

        while i < len(lines):
            line = lines[i]
            line_lower = line.lower().strip()

            # Check if we're after table unit headings (đơn vị tính) and before content
            # Look back up to 5 lines to find a table unit heading
            found_table_unit = False
            for j in range(max(0, i-5), i):
                if 'đơn vị tính' in lines[j].lower() and lines[j].strip().startswith('###'):
                    found_table_unit = True
                    break

            # If we found a table unit heading and this is a non-heading, non-table line
            if (found_table_unit and
                line_lower and
                not line.strip().startswith('#') and
                not line.strip().startswith('|') and
                len(line_lower) < 50):  # Short lines are typical of orphaned content

                # Look ahead to count similar lines (table content pattern)
                orphaned_count = 0
                temp_i = i
                while (temp_i < len(lines) and
                       not lines[temp_i].strip().startswith('#') and
                       not lines[temp_i].strip().startswith('|')):
                    if lines[temp_i].strip():
                        line_text = lines[temp_i].strip().lower()
                        # Normalize spaces for matching
                        normalized_text = ' '.join(line_text.split())
                        if (any(keyword in normalized_text for keyword in table_content_words) or
                            line_text.replace('.', '').replace(',', '').replace(' ', '').replace('000', '').isdigit()):
                            orphaned_count += 1
                        else:
                            break
                    temp_i += 1

                # If we found many orphaned table lines, skip them (they're likely duplicates)
                if orphaned_count >= 5:
                    # Skip all these orphaned lines
                    while (i < len(lines) and
                           not lines[i].strip().startswith('#') and
                           not lines[i].strip().startswith('|')):
                        if lines[i].strip():
                            line_text = lines[i].strip().lower()
                            # Normalize spaces for matching
                            normalized_text = ' '.join(line_text.split())
                            # Skip if it looks like table content
                            if (any(keyword in normalized_text for keyword in table_content_words) or
                                line_text.replace('.', '').replace(',', '').replace(' ', '').replace('000', '').isdigit()):
                                i += 1
                                continue
                        # Stop skipping if we hit something that's not table content
                        break
                    continue

            result.append(line)
            i += 1

        return '\n'.join(result)

    def _clean_markdown(self, markdown: str) -> str:
        """Clean up markdown formatting."""
        import re

        # Remove excessive blank lines (more than 2 consecutive)
        markdown = re.sub(r'\n{3,}', '\n\n', markdown)

        # Remove trailing spaces
        lines = [line.rstrip() for line in markdown.split('\n')]
        markdown = '\n'.join(lines)

        # Ensure proper spacing around headings
        markdown = re.sub(r'(\n#{1,6}\s+[^\n]+)\n', r'\1\n\n', markdown)

        return markdown.strip()

    def extract_document_metadata(self, markdown: str, source_metadata: Dict) -> Dict[str, Any]:
        """
        Extract metadata from the markdown content.

        Args:
            markdown: Markdown content
            source_metadata: Metadata from file

        Returns:
            Enhanced metadata dictionary
        """
        import re

        metadata = source_metadata.copy()

        # Extract title from first heading
        title_match = re.search(r'^#\s+(.+)', markdown, re.MULTILINE)
        if title_match and not metadata.get("title"):
            metadata["title"] = title_match.group(1).strip()

        # Count different elements
        metadata["heading_count"] = len(re.findall(r'^#{1,6}\s+', markdown, re.MULTILINE))
        metadata["table_count"] = len(re.findall(r'^\|.+\|$', markdown, re.MULTILINE))
        metadata["list_count"] = len(re.findall(r'^\s*[-*+]\s+', markdown, re.MULTILINE))

        # Estimate word count
        text_only = re.sub(r'[#*`\[\]|]', '', markdown)
        words = text_only.split()
        metadata["word_count"] = len(words)

        return metadata
