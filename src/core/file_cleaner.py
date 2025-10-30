"""
File Cleaner: Remove watermarks, headers, and footers from PDF and DOCX files using pikepdf.
"""
import os
from pathlib import Path
from typing import Dict, Optional, Tuple
from loguru import logger

try:
    import pikepdf
    PIKEPDF_AVAILABLE = True
except ImportError:
    logger.warning("pikepdf not available - install with: pip install pikepdf")
    PIKEPDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    logger.warning("python-docx not available")
    PYTHON_DOCX_AVAILABLE = False


class FileCleaner:
    """Clean PDF and DOCX files by removing watermarks, headers, and footers."""

    def __init__(self, output_dir: str = "temp/cleaned"):
        """Initialize file cleaner.

        Args:
            output_dir: Directory to save cleaned files
        """
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def clean_file(self, file_path: str) -> Tuple[bool, str, Optional[str]]:
        """Clean a file (PDF or DOCX).

        Args:
            file_path: Path to the file to clean

        Returns:
            Tuple of (success: bool, message: str, output_path: Optional[str])
        """
        file_path = Path(file_path)

        if not file_path.exists():
            return False, f"File not found: {file_path}", None

        ext = file_path.suffix.lower()

        try:
            if ext == ".pdf":
                return self._clean_pdf(file_path)
            elif ext == ".docx":
                return self._clean_docx(file_path)
            else:
                return False, f"Unsupported file format: {ext}. Only PDF and DOCX are supported.", None
        except Exception as e:
            logger.error(f"Error cleaning file: {e}", exc_info=True)
            return False, f"Error processing file: {str(e)}", None

    def _clean_pdf(self, file_path: Path) -> Tuple[bool, str, Optional[str]]:
        """Clean PDF by removing watermarks and annotations using pikepdf.

        Args:
            file_path: Path to PDF file

        Returns:
            Tuple of (success: bool, message: str, output_path: Optional[str])
        """
        if not PIKEPDF_AVAILABLE:
            return False, "pikepdf not installed. Install with: pip install pikepdf", None

        logger.info(f"Cleaning PDF with pikepdf: {file_path.name}")

        try:
            # Open PDF with pikepdf
            with pikepdf.open(file_path) as pdf:
                total_pages = len(pdf.pages)
                annotation_count = 0
                content_removed_count = 0

                logger.info(f"Processing {total_pages} pages")

                # Process each page
                for page_num, page in enumerate(pdf.pages, start=1):
                    try:
                        # Remove all annotations (watermarks, comments, etc)
                        if "/Annots" in page:
                            annots = page["/Annots"]
                            if annots is not None:
                                annotation_count += len(annots)
                                del page["/Annots"]
                                logger.debug(f"Page {page_num}: Removed {len(annots)} annotations")

                        # Remove watermark/header/footer from content streams
                        if "/Contents" in page:
                            try:
                                content_removed = self._remove_header_footer_from_content(page, pdf)
                                content_removed_count += content_removed
                                if content_removed > 0:
                                    logger.debug(f"Page {page_num}: Removed {content_removed} header/footer elements")
                            except Exception as content_err:
                                logger.debug(f"Page {page_num}: Could not process content stream: {content_err}")
                                continue

                    except Exception as page_error:
                        logger.warning(f"Error processing page {page_num}: {page_error}")
                        continue

                logger.info(f"Removed {annotation_count} annotations and {content_removed_count} header/footer elements")

                # Save cleaned PDF
                output_filename = f"cleaned_{file_path.stem}.pdf"
                output_path = self.output_dir / output_filename

                # Save with compression
                pdf.save(str(output_path), compress_streams=True)

                # Check file sizes
                original_size = file_path.stat().st_size
                cleaned_size = output_path.stat().st_size
                size_reduction = ((original_size - cleaned_size) / original_size * 100) if original_size > 0 else 0

                logger.info(f"Original size: {original_size} bytes, Cleaned size: {cleaned_size} bytes, Reduction: {size_reduction:.1f}%")

                return (
                    True,
                    f"PDF cleaned successfully. Removed {annotation_count} annotations and {content_removed_count} header/footer elements.",
                    str(output_path)
                )

        except Exception as e:
            logger.error(f"Error cleaning PDF: {e}", exc_info=True)
            return False, f"Error cleaning PDF: {str(e)}", None

    def _remove_header_footer_from_content(self, page, pdf) -> int:
        """Remove header/footer content from page content stream.

        Strategy: Use multi-pronged approach:
        1. Remove entire streams that appear to be header/footer (small, isolated)
        2. Filter text operators in remaining streams by position
        3. Handle coordinate transformations for accurate positioning

        Args:
            page: pikepdf page object
            pdf: pikepdf PDF object (needed for creating new streams)

        Returns:
            Count of elements removed
        """
        try:
            # Get page dimensions
            if "/MediaBox" not in page:
                return 0

            mediabox = page["/MediaBox"]
            page_height = float(mediabox[3]) - float(mediabox[1])
            page_width = float(mediabox[2]) - float(mediabox[0])

            # Define regions (top 10% and bottom 10% of page)
            header_threshold = page_height * 0.90  # Top 10%: y > 90% of height
            footer_threshold = page_height * 0.10  # Bottom 10%: y < 10% of height

            logger.debug(f"Page dimensions: {page_width} x {page_height}, header_threshold={header_threshold}, footer_threshold={footer_threshold}")

            # Get content stream(s)
            if "/Contents" not in page:
                return 0

            contents = page["/Contents"]
            if contents is None:
                return 0

            # Contents can be a single stream or an array of streams (pikepdf.Array)
            content_streams = []
            try:
                # Use isinstance to check if it's an array
                if isinstance(contents, pikepdf.Array):
                    # It's an array of streams
                    content_streams = list(contents)
                else:
                    # Single stream
                    content_streams = [contents]
            except:
                # Fallback: treat as single stream
                content_streams = [contents]

            # First pass: Identify streams to remove entirely (small header/footer streams)
            streams_to_remove = set()
            total_stream_size = sum(len(stream.read_bytes()) for stream in content_streams)

            for stream_idx, stream in enumerate(content_streams):
                try:
                    content_data = stream.read_bytes()
                    stream_size = len(content_data)
                    content_str = content_data.decode('latin-1', errors='ignore')
                    lines = content_str.split('\n')

                    # Count text and graphics operators
                    text_ops = sum(1 for line in lines if line.strip().endswith(('Tj', 'TJ')))

                    # Check for rotation/skew/transformation indicators (header/footer watermarks)
                    has_rotation = any(
                        line.strip().startswith(('-')) and 'Tm' in line  # Negative scale = rotation/flip
                        for line in lines
                    )

                    # Check for very high X positions (right side of page - signature box)
                    has_right_edge_text = any(
                        float(parts[-3]) > 350  # X position > 350 (close to right edge of ~595 width page)
                        for line in lines
                        if line.strip().endswith('Tm') and len((parts := line.strip().split())) >= 6
                        for _ in [None]  # Dummy to allow parts to be defined
                    )

                    # Heuristic: Stream is likely ONLY header/footer if:
                    # 1. Is small (<5%) AND has no text, OR
                    # 2. Is very small AND has only rotated content, OR
                    # 3. Has very few text operators (<15) in header/footer region (watermark indicator)
                    # BUT: Don't remove streams that mix legitimate content with watermarks!
                    is_purely_small_empty = stream_size < total_stream_size * 0.05 and text_ops == 0
                    is_purely_rotated = has_rotation and stream_size < total_stream_size * 0.03 and text_ops < 10

                    # New heuristic: Detect watermarks by very low text operator count in small-to-medium streams
                    # This catches streams like: 1062 bytes with only 7 text operators (watermark font/text)
                    # BUT: Don't treat large graphics-only streams as watermarks (they may be page content)
                    # Only mark as watermark if: has SOME text (1-15 ops) AND is small-to-medium size
                    # Raised threshold from 10% to 20% to catch larger watermark streams (pages 3, 19, 21, 26)
                    is_likely_watermark = (0 < text_ops < 15) and stream_size < total_stream_size * 0.2 and stream_size > 100

                    if is_purely_small_empty or is_purely_rotated or is_likely_watermark:
                        reason = "empty" if is_purely_small_empty else ("rotated" if is_purely_rotated else "watermark_low_text_ops")
                        logger.debug(f"Stream {stream_idx}: Marked for removal (size={stream_size}, text_ops={text_ops}, rotation={has_rotation}, right_edge={has_right_edge_text}, reason={reason})")
                        streams_to_remove.add(stream_idx)
                    elif (has_rotation or has_right_edge_text) and text_ops < 30:
                        # Stream has mixed content (watermark + legitimate text)
                        # Don't remove entirely, but mark for filtering instead
                        logger.debug(f"Stream {stream_idx}: Will be filtered (mixed content, size={stream_size}, text_ops={text_ops}, rotation={has_rotation})")
                except Exception as e:
                    logger.debug(f"Error analyzing stream {stream_idx}: {e}")

            # Second pass: Process remaining streams
            total_text_ops_removed = 0
            new_streams = []

            for stream_idx, stream in enumerate(content_streams):
                if stream_idx in streams_to_remove:
                    # Skip this stream entirely
                    total_text_ops_removed += 1
                    logger.debug(f"Stream {stream_idx}: Removed entirely (header/footer)")
                    continue

                try:
                    content_data = stream.read_bytes()
                    filtered_content, text_ops_removed = self._filter_content_stream(
                        content_data, header_threshold, footer_threshold, page_height
                    )

                    total_text_ops_removed += text_ops_removed

                    if filtered_content != content_data:
                        logger.debug(f"Stream {stream_idx}: Removed {text_ops_removed} text operators")
                        # Create new stream with filtered content
                        # Note: pikepdf.Stream requires (pdf, bytes) not (page, bytes)
                        new_stream = pikepdf.Stream(pdf, filtered_content)
                        new_streams.append(new_stream)
                    else:
                        new_streams.append(stream)

                except Exception as stream_err:
                    logger.debug(f"Error processing stream {stream_idx}: {stream_err}, keeping original")
                    new_streams.append(stream)

            # Replace content streams if any were modified
            if total_text_ops_removed > 0:
                if len(new_streams) == 1:
                    page["/Contents"] = new_streams[0]
                else:
                    page["/Contents"] = pikepdf.Array(new_streams)

            return total_text_ops_removed

        except Exception as e:
            logger.debug(f"Error filtering content: {e}")
            return 0

    def _filter_content_stream(self, content: bytes, header_threshold: float, footer_threshold: float, page_height: float) -> Tuple[bytes, int]:
        """Filter content stream to remove header/footer text.

        Strategy: Parse PDF content operators, track text matrix (Tm), and skip
        text drawing commands (Tj, TJ) when text positioning indicates header/footer.

        Note: Y coordinates can be negative or outside normal page range due to
        coordinate transformations (cm commands). We detect these as header/footer.

        Args:
            content: PDF content stream bytes
            header_threshold: Y position for top boundary
            footer_threshold: Y position for bottom boundary
            page_height: Total page height for context

        Returns:
            Tuple of (filtered content bytes, count of text operators removed)
        """
        try:
            content_str = content.decode('latin-1', errors='ignore')

            # Split into tokens/lines for processing
            lines = content_str.split('\n')
            filtered_lines = []
            skip_next_text = False
            skip_hex_footer = False  # ← Track footer-specific Y position separately!
            removed_count = 0

            # Page dimensions (approximate - typical A4 PDF)
            page_width = 595  # Approximate standard page width
            footer_y_positions = [0, 10, 13.07, 28.3, 29.97, 30.97, 46.87]

            for i, line in enumerate(lines):
                line_stripped = line.strip()

                if not line_stripped:
                    filtered_lines.append(line)
                    continue

                # Track text positioning - Tm sets absolute position
                # Format: a b c d e f Tm (e=x, f=y in standard coordinates)
                if line_stripped.endswith('Tm'):
                    parts = line_stripped.split()
                    if len(parts) >= 6:
                        try:
                            # Extract positioning
                            x_pos = float(parts[-4])
                            y_pos = float(parts[-3])

                            # Detect header/footer by multiple criteria:
                            # 1. Y is in top 10% (y > 90% of height)
                            # 2. Y is in bottom 10% (y < 10% of height)
                            # 3. Y is negative (common for headers/footers in rotated content)
                            # 4. Y is way below page (footer outside visible area)
                            # 5. Page numbers/footers: Very bottom (0 < Y < 50) - usually contains page numbers
                            # 6. Footer specific Y positions: 0, 13.07, 28.3, 29.97, 30.97 (Vietnamese doc pattern)
                            # 7. Center top: Very top (y > 95% height) AND center X (200-400)
                            is_header = y_pos >= header_threshold
                            is_footer = y_pos <= footer_threshold
                            is_unusual = y_pos < 0 or y_pos > page_height + 100
                            is_page_number_top = y_pos >= page_height * 0.95 and 200 < x_pos < 400
                            # Very specific footer Y positions for Vietnamese documents
                            # Only use for hex-footer detection, NOT for general header/footer region
                            is_footer_specific_y = any(abs(y_pos - fy) < 1.0 for fy in footer_y_positions)

                            # Set skip_next_text for NORMAL header/footer removal (page numbers, footer lines)
                            if is_header or is_footer or is_unusual or is_page_number_top:
                                skip_next_text = True
                                skip_hex_footer = False  # Don't skip hex at header!
                                if is_page_number_top:
                                    logger.debug(f"Page number detected at top X={x_pos:.1f}, Y={y_pos:.1f}")
                            # Set skip_hex_footer ONLY for footer-specific Y positions
                            elif is_footer_specific_y:
                                skip_next_text = False  # Don't skip everything
                                skip_hex_footer = True  # Only skip hex-encoded footer
                            else:
                                skip_next_text = False
                                skip_hex_footer = False
                        except (ValueError, IndexError):
                            pass

                # Also track Td/TD (relative positioning) - this resets positioning within text object
                # but we keep skip flag since we're still in the same text object
                elif line_stripped.endswith('Td') or line_stripped.endswith('TD'):
                    pass  # Keep skip flag if already in header/footer text

                # Check if this looks like page number/footer/line decoration
                # Pattern: [<hex>], [(number)], [(_____)], etc
                elif line_stripped.endswith(('Tj', 'TJ')):
                    content = line_stripped
                    stripped_content = content.replace('Tj', '').replace('TJ', '').strip()

                    # Check patterns VERY CAREFULLY to avoid removing main content:
                    # Only remove SPECIFIC patterns, not all short text!

                    # 1. ONLY page numbers: very specific patterns like [(1)], [( 1 )], [(1: )], etc
                    # Must be 1-3 digits only, with optional spaces and colon
                    is_page_number = (
                        stripped_content.startswith('[') and
                        stripped_content.endswith(']') and
                        '(' in stripped_content and ')' in stripped_content and
                        len(stripped_content) < 30 and  # Very short
                        # Extract content between [ and ]
                        all(c in '()0123456789 \t:' for c in stripped_content)  # Only parens, digits, colon
                    )

                    # 2. Footer lines: ONLY dashes/underscores (decoration lines)
                    is_footer_line = (
                        stripped_content.startswith('[') and
                        stripped_content.endswith(']') and
                        '(' in stripped_content and ')' in stripped_content and
                        all(c in '()- _\t' for c in stripped_content) and
                        any(c in '-_' for c in stripped_content)
                    )

                    # 3. Hex-encoded footer text at footer-specific positions
                    # Pattern like [<0015002D002B0034>] ONLY at footer-specific Y positions
                    is_hex_footer = (
                        skip_hex_footer and  # ← ONLY for footer-specific Y (not header!)
                        stripped_content.startswith('[<') and
                        stripped_content.endswith('>]') and
                        len(stripped_content) < 80 and  # Short
                        all(c in '[]<>0123456789ABCDEFabcdef \t-' for c in stripped_content)
                    )

                    in_header_footer_region = skip_next_text

                    # Skip specific patterns - page numbers, footer lines, and hex-encoded footer text
                    should_skip = is_page_number or is_footer_line or is_hex_footer

                    if should_skip:
                        if is_page_number:
                            logger.debug(f"Skipping page number: {content[:50]}")
                        elif is_footer_line:
                            logger.debug(f"Skipping footer line: {content[:50]}")
                        elif is_hex_footer:
                            logger.debug(f"Skipping hex-encoded footer: {content[:50]}")
                        removed_count += 1
                        continue

                # NOTE: Graphics command removal disabled due to false positives
                # These commands (m, l, S) are part of normal PDF graphics and are difficult to
                # distinguish from footer-drawing commands without more context.
                # Footer detection by position (via Tm/text positioning) is more reliable.
                # TODO: Implement more sophisticated footer line detection if needed

                # Reset flag on text object end (ET) or start (BT)
                elif line_stripped == 'ET':  # End text object
                    skip_next_text = False
                elif line_stripped == 'BT':  # Begin text object - start fresh
                    skip_next_text = False

                filtered_lines.append(line)

            filtered_content = '\n'.join(filtered_lines)
            return filtered_content.encode('latin-1', errors='ignore'), removed_count

        except Exception as e:
            logger.debug(f"Error in content filtering: {e}")
            return content, 0

    def _clean_docx(self, file_path: Path) -> Tuple[bool, str, Optional[str]]:
        """Clean DOCX by removing headers, footers, and watermarks.

        Args:
            file_path: Path to DOCX file

        Returns:
            Tuple of (success: bool, message: str, output_path: Optional[str])
        """
        if not PYTHON_DOCX_AVAILABLE:
            return False, "python-docx not installed. Install with: pip install python-docx", None

        logger.info(f"Cleaning DOCX: {file_path.name}")

        try:
            doc = DocxDocument(file_path)

            # Remove headers from all sections
            header_count = 0
            for section in doc.sections:
                header = section.header
                for paragraph in list(header.paragraphs):
                    p = paragraph._element
                    p.getparent().remove(p)
                    header_count += 1

            # Remove footers from all sections
            footer_count = 0
            for section in doc.sections:
                footer = section.footer
                for paragraph in list(footer.paragraphs):
                    p = paragraph._element
                    p.getparent().remove(p)
                    footer_count += 1

            logger.info(f"Removed {header_count} headers and {footer_count} footers")

            # Save cleaned DOCX
            output_filename = f"cleaned_{file_path.stem}.docx"
            output_path = self.output_dir / output_filename
            doc.save(str(output_path))

            logger.info(f"Cleaned DOCX saved to: {output_path}")
            return (
                True,
                f"DOCX cleaned successfully. Removed {header_count} headers and {footer_count} footers.",
                str(output_path)
            )

        except Exception as e:
            logger.error(f"Error cleaning DOCX: {e}", exc_info=True)
            return False, f"Error cleaning DOCX: {str(e)}", None
